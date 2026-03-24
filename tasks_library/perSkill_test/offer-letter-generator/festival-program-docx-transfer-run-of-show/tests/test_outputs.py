import json
import re
from pathlib import Path

import pytest
from docx import Document

OUTPUT_FILE = "/root/festival_run_of_show_final.docx"
DATA_FILE = "/root/festival_run_of_show_data.json"


@pytest.fixture(scope="module")
def festival_data():
    with open(DATA_FILE, encoding="utf-8") as f:
        return json.load(f)


@pytest.fixture(scope="module")
def output_doc():
    output_path = Path(OUTPUT_FILE)
    assert output_path.exists(), f"Output file not found: {OUTPUT_FILE}"
    return Document(OUTPUT_FILE)


def collect_all_text(doc):
    parts = []

    def walk_container(container):
        for paragraph in container.paragraphs:
            parts.append(paragraph.text)
        for table in container.tables:
            walk_table(table)

    def walk_table(table):
        for row in table.rows:
            for cell in row.cells:
                walk_container(cell)

    walk_container(doc)

    for section in doc.sections:
        walk_container(section.header)
        walk_container(section.footer)

    return "\n".join(parts)


def iter_tables(container):
    for table in container.tables:
        yield table
        for row in table.rows:
            for cell in row.cells:
                yield from iter_tables(cell)


def find_table_by_header(doc, expected_header):
    for table in iter_tables(doc):
        if table.rows:
            header = [cell.text.strip() for cell in table.rows[0].cells]
            if header == expected_header:
                return table
    return None


def test_no_placeholder_markers_left(output_doc):
    all_text = collect_all_text(output_doc)
    assert not re.findall(r"\{\{[A-Z0-9_]+\}\}", all_text), all_text


def test_footer_line_replaced(output_doc, festival_data):
    footer_lines = [
        paragraph.text.strip()
        for section in output_doc.sections
        for paragraph in section.footer.paragraphs
        if paragraph.text.strip()
    ]
    assert footer_lines == [
        (
            f"Run of Show {festival_data['RUN_OF_SHOW_ID']} | "
            f"{festival_data['DOC_VERSION']} | "
            f"Ops Hotline {festival_data['OPS_HOTLINE']}"
        )
    ]


def test_schedule_table_updated(output_doc, festival_data):
    schedule_table = find_table_by_header(output_doc, ["Time", "Cue", "Lead", "Notes"])
    assert schedule_table is not None, "Run-of-show schedule table not found"

    actual_rows = [
        [cell.text.strip() for cell in row.cells]
        for row in schedule_table.rows
    ]
    expected_rows = [
        ["Time", "Cue", "Lead", "Notes"],
        [
            festival_data["SLOT_1_TIME"],
            festival_data["SLOT_1_CUE"],
            festival_data["SLOT_1_LEAD"],
            festival_data["SLOT_1_NOTES"],
        ],
        [
            festival_data["SLOT_2_TIME"],
            festival_data["SLOT_2_CUE"],
            festival_data["SLOT_2_LEAD"],
            festival_data["SLOT_2_NOTES"],
        ],
        [
            festival_data["SLOT_3_TIME"],
            festival_data["SLOT_3_CUE"],
            festival_data["SLOT_3_LEAD"],
            festival_data["SLOT_3_NOTES"],
        ],
        [
            festival_data["SLOT_4_TIME"],
            festival_data["SLOT_4_CUE"],
            festival_data["SLOT_4_LEAD"],
            festival_data["SLOT_4_NOTES"],
        ],
        [
            festival_data["SLOT_5_TIME"],
            festival_data["SLOT_5_CUE"],
            festival_data["SLOT_5_LEAD"],
            festival_data["SLOT_5_NOTES"],
        ],
    ]

    assert actual_rows == expected_rows


def test_nested_stage_prep_table_updated(output_doc, festival_data):
    stage_prep_table = find_table_by_header(output_doc, ["Area", "Ready Item", "Owner"])
    assert stage_prep_table is not None, "Stage prep table not found"

    actual_rows = [
        [cell.text.strip() for cell in row.cells]
        for row in stage_prep_table.rows
    ]
    expected_rows = [
        ["Area", "Ready Item", "Owner"],
        [
            festival_data["PREP_1_AREA"],
            festival_data["PREP_1_ITEM"],
            festival_data["PREP_1_OWNER"],
        ],
        [
            festival_data["PREP_2_AREA"],
            festival_data["PREP_2_ITEM"],
            festival_data["PREP_2_OWNER"],
        ],
        [
            festival_data["PREP_3_AREA"],
            festival_data["PREP_3_ITEM"],
            festival_data["PREP_3_OWNER"],
        ],
        [
            festival_data["PREP_4_AREA"],
            festival_data["PREP_4_ITEM"],
            festival_data["PREP_4_OWNER"],
        ],
    ]

    assert actual_rows == expected_rows


def test_contact_block_present(output_doc, festival_data):
    body_text = "\n".join(paragraph.text for paragraph in output_doc.paragraphs)
    expected_lines = [
        festival_data["FESTIVAL_NAME"],
        festival_data["EVENT_DATE"],
        festival_data["VENUE_NAME"],
        festival_data["STAGE_MANAGER"],
        festival_data["ARTIST_LIAISON"],
        festival_data["AUDIO_LEAD"],
        festival_data["LIGHTING_LEAD"],
    ]

    for value in expected_lines:
        assert value in body_text, f"Expected body value missing: {value}"


def test_rain_plan_clause_kept_without_markers(output_doc, festival_data):
    all_text = collect_all_text(output_doc)
    expected_sentence = (
        f"If weather escalation is triggered by {festival_data['RAIN_TRIGGER']}, "
        f"crew and artists move to {festival_data['RAIN_MEETING_POINT']} and await updates on "
        f"{festival_data['RAIN_UPDATE_CHANNEL']} from {festival_data['RAIN_LEAD']}."
    )

    assert "{{IF_RAIN_PLAN}}" not in all_text
    assert "{{END_IF_RAIN_PLAN}}" not in all_text
    assert expected_sentence in all_text
