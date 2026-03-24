import json
import re
from pathlib import Path

import pytest
from docx import Document

OUTPUT_FILE = "/root/discharge_packet_final.docx"
DATA_FILE = "/root/patient_discharge_data.json"


@pytest.fixture(scope="module")
def discharge_data():
    with open(DATA_FILE, encoding="utf-8") as f:
        return json.load(f)


@pytest.fixture(scope="module")
def output_doc():
    output_path = Path(OUTPUT_FILE)
    assert output_path.exists(), f"Output file not found: {OUTPUT_FILE}"
    return Document(OUTPUT_FILE)


def collect_all_text(doc):
    parts = []

    def walk_container(container):
        for paragraph in container.paragraphs:
            parts.append(paragraph.text)
        for table in container.tables:
            walk_table(table)

    def walk_table(table):
        for row in table.rows:
            for cell in row.cells:
                walk_container(cell)

    walk_container(doc)

    for section in doc.sections:
        walk_container(section.header)
        walk_container(section.footer)

    return "\n".join(parts)


def iter_tables(container):
    for table in container.tables:
        yield table
        for row in table.rows:
            for cell in row.cells:
                yield from iter_tables(cell)


def find_table_by_header(doc, expected_header):
    for table in iter_tables(doc):
        if table.rows:
            header = [cell.text.strip() for cell in table.rows[0].cells]
            if header == expected_header:
                return table
    for section in doc.sections:
        for table in iter_tables(section.header):
            if table.rows and [cell.text.strip() for cell in table.rows[0].cells] == expected_header:
                return table
        for table in iter_tables(section.footer):
            if table.rows and [cell.text.strip() for cell in table.rows[0].cells] == expected_header:
                return table
    return None


def test_no_placeholder_markers_left(output_doc):
    all_text = collect_all_text(output_doc)
    assert not re.findall(r"\{\{[A-Z_]+\}\}", all_text), all_text


def test_header_identity_replaced(output_doc, discharge_data):
    header_text = "\n".join(
        paragraph.text for section in output_doc.sections for paragraph in section.header.paragraphs
    )
    assert discharge_data["PATIENT_FULL_NAME"] in header_text
    assert discharge_data["MEDICAL_RECORD_NUMBER"] in header_text
    assert discharge_data["PACKET_ID"] in header_text


def test_summary_table_values(output_doc, discharge_data):
    summary_table = find_table_by_header(output_doc, ["Patient Name", discharge_data["PATIENT_FULL_NAME"]])
    assert summary_table is not None, "Summary table row for patient name not found"

    summary_pairs = {
        row.cells[0].text.strip(): row.cells[1].text.strip()
        for row in summary_table.rows
        if len(row.cells) == 2
    }

    expected_pairs = {
        "Patient Name": discharge_data["PATIENT_FULL_NAME"],
        "Date of Birth": discharge_data["PATIENT_DOB"],
        "Primary Diagnosis": discharge_data["PRIMARY_DIAGNOSIS"],
        "Disposition": discharge_data["DISPOSITION"],
        "Attending Clinician": discharge_data["ATTENDING_CLINICIAN"],
    }

    assert summary_pairs == expected_pairs


def test_medication_table_expanded(output_doc, discharge_data):
    medication_table = find_table_by_header(output_doc, ["Medication", "Dose", "Schedule", "Purpose"])
    assert medication_table is not None, "Medication table not found"
    assert len(medication_table.rows) == 1 + len(discharge_data["MEDICATIONS"])

    actual_rows = [
        [cell.text.strip() for cell in row.cells]
        for row in medication_table.rows[1:]
    ]
    expected_rows = [
        [item["NAME"], item["DOSE"], item["SCHEDULE"], item["PURPOSE"]]
        for item in discharge_data["MEDICATIONS"]
    ]

    assert actual_rows == expected_rows


def test_followup_table_expanded(output_doc, discharge_data):
    followup_table = find_table_by_header(output_doc, ["Clinic", "Date/Time", "Location", "Notes"])
    assert followup_table is not None, "Follow-up table not found"
    assert len(followup_table.rows) == 1 + len(discharge_data["FOLLOW_UP_APPOINTMENTS"])

    actual_rows = [
        [cell.text.strip() for cell in row.cells]
        for row in followup_table.rows[1:]
    ]
    expected_rows = [
        [item["CLINIC"], item["WHEN"], item["LOCATION"], item["NOTES"]]
        for item in discharge_data["FOLLOW_UP_APPOINTMENTS"]
    ]

    assert actual_rows == expected_rows


def test_remote_followup_clause_kept_without_markers(output_doc, discharge_data):
    all_text = collect_all_text(output_doc)
    expected_sentence = (
        f"Remote follow-up is required {discharge_data['REMOTE_FOLLOWUP_WINDOW']}. "
        f"Please connect through {discharge_data['REMOTE_FOLLOWUP_PLATFORM']} "
        f"or call {discharge_data['REMOTE_FOLLOWUP_PHONE']}."
    )

    assert "{{IF_REMOTE_FOLLOWUP}}" not in all_text
    assert "{{END_IF_REMOTE_FOLLOWUP}}" not in all_text
    assert expected_sentence in all_text
