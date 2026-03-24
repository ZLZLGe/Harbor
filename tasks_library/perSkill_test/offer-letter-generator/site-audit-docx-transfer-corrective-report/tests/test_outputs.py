import json
import re
from pathlib import Path

import pytest
from docx import Document

OUTPUT_FILE = "/root/site_audit_report_final.docx"
DATA_FILE = "/root/site_audit_findings.json"


@pytest.fixture(scope="module")
def audit_data():
    with open(DATA_FILE, encoding="utf-8") as f:
        return json.load(f)


@pytest.fixture(scope="module")
def output_doc():
    output_path = Path(OUTPUT_FILE)
    assert output_path.exists(), f"Output file not found: {OUTPUT_FILE}"
    return Document(OUTPUT_FILE)


def collect_all_text(doc):
    parts = []

    def walk_container(container):
        for paragraph in container.paragraphs:
            parts.append(paragraph.text)
        for table in container.tables:
            walk_table(table)

    def walk_table(table):
        for row in table.rows:
            for cell in row.cells:
                walk_container(cell)

    walk_container(doc)

    for section in doc.sections:
        walk_container(section.header)
        walk_container(section.footer)

    return "\n".join(parts)


def iter_tables(container):
    for table in container.tables:
        yield table
        for row in table.rows:
            for cell in row.cells:
                yield from iter_tables(cell)


def find_matrix_table(doc):
    expected_header = ["Control Area", "Result", "Severity", "Evidence"]
    for table in iter_tables(doc):
        if table.rows:
            header = [cell.text.strip() for cell in table.rows[0].cells]
            if header == expected_header:
                return table
    return None


def find_summary_table(doc):
    for table in doc.tables:
        if not table.rows:
            continue
        first_row = [cell.text.strip() for cell in table.rows[0].cells]
        if first_row == ["Report Summary", "Value"]:
            return table
    return None


def test_no_placeholder_markers_left(output_doc):
    all_text = collect_all_text(output_doc)
    assert not re.findall(r"\{\{[A-Z_]+\}\}", all_text), all_text


def test_footer_replaced(output_doc, audit_data):
    footer_text = "\n".join(
        paragraph.text for section in output_doc.sections for paragraph in section.footer.paragraphs
    )
    assert audit_data["REPORT_ID"] in footer_text
    assert audit_data["SITE_CODE"] in footer_text
    assert audit_data["CLIENT_NAME"] in footer_text


def test_body_values_present(output_doc, audit_data):
    body_text = "\n".join(paragraph.text for paragraph in output_doc.paragraphs)
    required_values = [
        audit_data["SITE_NAME"],
        audit_data["AUDIT_DATE"],
        audit_data["CLIENT_NAME"],
        audit_data["LEAD_AUDITOR"],
        audit_data["SECOND_AUDITOR"],
        audit_data["AUDIT_SCOPE"],
    ]

    for value in required_values:
        assert value in body_text, f"Expected value not found in body: {value}"


def test_summary_table_updated(output_doc, audit_data):
    summary_table = find_summary_table(output_doc)
    assert summary_table is not None, "Summary table not found"

    actual_rows = [
        [cell.text.strip() for cell in row.cells]
        for row in summary_table.rows
    ]
    expected_rows = [
        ["Report Summary", "Value"],
        ["Overall Rating", audit_data["OVERALL_RATING"]],
        ["Critical Findings", audit_data["CRITICAL_FINDINGS"]],
        ["Prepared For", audit_data["CLIENT_NAME"]],
    ]

    assert actual_rows == expected_rows


def test_nested_compliance_matrix_updated(output_doc, audit_data):
    matrix_table = find_matrix_table(output_doc)
    assert matrix_table is not None, "Compliance matrix table not found"

    actual_rows = [
        [cell.text.strip() for cell in row.cells]
        for row in matrix_table.rows
    ]
    expected_rows = [
        ["Control Area", "Result", "Severity", "Evidence"],
        [
            audit_data["CONTROL_1_AREA"],
            audit_data["CONTROL_1_STATUS"],
            audit_data["CONTROL_1_SEVERITY"],
            audit_data["CONTROL_1_EVIDENCE"],
        ],
        [
            audit_data["CONTROL_2_AREA"],
            audit_data["CONTROL_2_STATUS"],
            audit_data["CONTROL_2_SEVERITY"],
            audit_data["CONTROL_2_EVIDENCE"],
        ],
        [
            audit_data["CONTROL_3_AREA"],
            audit_data["CONTROL_3_STATUS"],
            audit_data["CONTROL_3_SEVERITY"],
            audit_data["CONTROL_3_EVIDENCE"],
        ],
        [
            audit_data["CONTROL_4_AREA"],
            audit_data["CONTROL_4_STATUS"],
            audit_data["CONTROL_4_SEVERITY"],
            audit_data["CONTROL_4_EVIDENCE"],
        ],
    ]

    assert actual_rows == expected_rows


def test_corrective_action_clause_removed(output_doc, audit_data):
    all_text = collect_all_text(output_doc)

    assert "{{IF_CRITICAL_FINDINGS}}" not in all_text
    assert "{{END_IF_CRITICAL_FINDINGS}}" not in all_text
    assert "Corrective action is required within" not in all_text
    assert audit_data["FOLLOW_UP_WINDOW"] not in all_text
    assert audit_data["ESCALATION_CONTACT"] not in all_text
