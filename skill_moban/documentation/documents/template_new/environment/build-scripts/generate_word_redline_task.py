from __future__ import annotations

import json
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
REVIEW_NS = "urn:northharbor:review-manifest"

ET.register_namespace("w", W_NS)
ET.register_namespace("rv", REVIEW_NS)


def qn(namespace: str, name: str) -> str:
    return f"{{{namespace}}}{name}"


def run_el(text: str, *, bold: bool = False, style: str | None = None) -> ET.Element:
    run = ET.Element(qn(W_NS, "r"))
    if bold or style:
        rpr = ET.SubElement(run, qn(W_NS, "rPr"))
        if bold:
            ET.SubElement(rpr, qn(W_NS, "b"))
        if style:
            ET.SubElement(rpr, qn(W_NS, "rStyle"), {qn(W_NS, "val"): style})
    text_node = ET.SubElement(run, qn(W_NS, "t"))
    if text.startswith(" ") or text.endswith(" "):
        text_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_node.text = text
    return run


def deletion_el(text: str, *, revision_id: int, author: str) -> ET.Element:
    deletion = ET.Element(
        qn(W_NS, "del"),
        {
            qn(W_NS, "id"): str(revision_id),
            qn(W_NS, "author"): author,
            qn(W_NS, "date"): "2026-04-18T09:00:00Z",
        },
    )
    run = ET.SubElement(deletion, qn(W_NS, "r"))
    text_node = ET.SubElement(run, qn(W_NS, "delText"))
    if text.startswith(" ") or text.endswith(" "):
        text_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_node.text = text
    return deletion


def insertion_el(text: str, *, revision_id: int, author: str) -> ET.Element:
    insertion = ET.Element(
        qn(W_NS, "ins"),
        {
            qn(W_NS, "id"): str(revision_id),
            qn(W_NS, "author"): author,
            qn(W_NS, "date"): "2026-04-18T09:05:00Z",
        },
    )
    insertion.append(run_el(text))
    return insertion


def comment_reference_run(comment_id: int) -> ET.Element:
    run = ET.Element(qn(W_NS, "r"))
    rpr = ET.SubElement(run, qn(W_NS, "rPr"))
    ET.SubElement(rpr, qn(W_NS, "rStyle"), {qn(W_NS, "val"): "CommentReference"})
    ET.SubElement(run, qn(W_NS, "commentReference"), {qn(W_NS, "id"): str(comment_id)})
    return run


def footnote_reference_run(footnote_id: int) -> ET.Element:
    run = ET.Element(qn(W_NS, "r"))
    rpr = ET.SubElement(run, qn(W_NS, "rPr"))
    ET.SubElement(rpr, qn(W_NS, "rStyle"), {qn(W_NS, "val"): "FootnoteReference"})
    ET.SubElement(run, qn(W_NS, "footnoteReference"), {qn(W_NS, "id"): str(footnote_id)})
    return run


def wrap_with_comment(comment_id: int, content: list[ET.Element]) -> list[ET.Element]:
    wrapped: list[ET.Element] = [ET.Element(qn(W_NS, "commentRangeStart"), {qn(W_NS, "id"): str(comment_id)})]
    wrapped.extend(content)
    wrapped.append(ET.Element(qn(W_NS, "commentRangeEnd"), {qn(W_NS, "id"): str(comment_id)}))
    wrapped.append(comment_reference_run(comment_id))
    return wrapped


def replace_paragraph_children(paragraph: ET.Element, new_children: list[ET.Element]) -> None:
    ppr = paragraph.find(qn(W_NS, "pPr"))
    paragraph.clear()
    if ppr is not None:
        paragraph.append(ppr)
    for child in new_children:
        paragraph.append(child)


REVIEW_ITEMS = [
    {
        "comment_id": 0,
        "review_ref": "RV-TERM-01",
        "decision_key": "TERM_EXTENSION",
        "part": "word/document.xml",
        "reviewer_note": "Accept the supplier redline that extends the renewed term through April 30, 2027.",
    },
    {
        "comment_id": 1,
        "review_ref": "RV-NOTICE-02",
        "decision_key": "NOTICE_PERIOD",
        "part": "word/document.xml",
        "reviewer_note": "Accept the legal redline that increases the non-renewal notice period to 45 days.",
    },
    {
        "comment_id": 2,
        "review_ref": "RV-SEC-03",
        "decision_key": "SECURITY_REVIEW",
        "part": "word/document.xml",
        "reviewer_note": "Accept the inserted requirement for a current security review packet before the renewal takes effect.",
    },
    {
        "comment_id": 3,
        "review_ref": "RV-USAGE-04",
        "decision_key": "USAGE_LOG_RETENTION",
        "part": "word/document.xml",
        "reviewer_note": "Accept the updated usage-log retention period for the renewal term.",
    },
    {
        "comment_id": 4,
        "review_ref": "RV-AUDIT-05",
        "decision_key": "AUDIT_REPORT_WINDOW",
        "part": "word/footnotes.xml",
        "reviewer_note": "Reject the shorter audit reporting window and keep the existing timing.",
    },
    {
        "comment_id": 5,
        "review_ref": "RV-PRICE-06",
        "decision_key": "PRICING_HOLD",
        "part": "word/document.xml",
        "reviewer_note": "Reject the vendor pricing increase and keep the existing annual fee unchanged.",
    },
    {
        "comment_id": 6,
        "review_ref": "RV-LICENSE-07",
        "decision_key": "LICENSE_COUNT",
        "part": "word/document.xml",
        "reviewer_note": "Accept the increase in licensed users for the renewal term.",
    },
    {
        "comment_id": 7,
        "review_ref": "RV-CREDIT-08",
        "decision_key": "SERVICE_CREDIT_CAP",
        "part": "word/document.xml",
        "reviewer_note": "Reject the proposed increase to the service credit cap and keep the existing amount.",
    },
]


def make_comments_xml() -> bytes:
    comments = ET.Element(qn(W_NS, "comments"))
    for item in REVIEW_ITEMS:
        comment = ET.SubElement(
            comments,
            qn(W_NS, "comment"),
            {
                qn(W_NS, "id"): str(item["comment_id"]),
                qn(W_NS, "author"): "Legal Ops",
                qn(W_NS, "initials"): "LO",
                qn(W_NS, "date"): "2026-04-18T09:10:00Z",
            },
        )
        for line in [f"Review Ref: {item['review_ref']}", str(item["reviewer_note"])]:
            para = ET.SubElement(comment, qn(W_NS, "p"))
            para.append(run_el(line))
    return ET.tostring(comments, encoding="utf-8", xml_declaration=True)


def make_review_manifest_xml() -> bytes:
    root = ET.Element(qn(REVIEW_NS, "reviewManifest"))
    root.set("documentId", "vendor-addendum-redline")
    for item in REVIEW_ITEMS:
        ET.SubElement(
            root,
            qn(REVIEW_NS, "item"),
            {
                "reviewRef": str(item["review_ref"]),
                "decisionKey": str(item["decision_key"]),
                "part": str(item["part"]),
                "commentId": str(item["comment_id"]),
                "status": "pending",
            },
        )
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def make_review_manifest_props_xml() -> bytes:
    root = ET.Element(qn("http://schemas.openxmlformats.org/officeDocument/2006/customXml", "datastoreItem"))
    root.set(qn("http://schemas.openxmlformats.org/officeDocument/2006/customXml", "itemID"), "{38E790A0-64A4-4C68-BA4C-C6B27F61B97A}")
    schema_refs = ET.SubElement(root, qn("http://schemas.openxmlformats.org/officeDocument/2006/customXml", "schemaRefs"))
    ET.SubElement(
        schema_refs,
        qn("http://schemas.openxmlformats.org/officeDocument/2006/customXml", "schemaRef"),
        {qn("http://schemas.openxmlformats.org/officeDocument/2006/customXml", "uri"): REVIEW_NS},
    )
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def make_review_manifest_rels_xml() -> bytes:
    root = ET.Element(qn(PKG_REL_NS, "Relationships"))
    ET.SubElement(
        root,
        qn(PKG_REL_NS, "Relationship"),
        {
            "Id": "rId1",
            "Type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXmlProps",
            "Target": "itemProps1.xml",
        },
    )
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def make_footnotes_xml() -> bytes:
    root = ET.Element(qn(W_NS, "footnotes"))

    separator = ET.SubElement(root, qn(W_NS, "footnote"), {qn(W_NS, "type"): "separator", qn(W_NS, "id"): "-1"})
    sep_para = ET.SubElement(separator, qn(W_NS, "p"))
    sep_run = ET.SubElement(sep_para, qn(W_NS, "r"))
    ET.SubElement(sep_run, qn(W_NS, "separator"))

    continuation = ET.SubElement(
        root,
        qn(W_NS, "footnote"),
        {qn(W_NS, "type"): "continuationSeparator", qn(W_NS, "id"): "0"},
    )
    cont_para = ET.SubElement(continuation, qn(W_NS, "p"))
    cont_run = ET.SubElement(cont_para, qn(W_NS, "r"))
    ET.SubElement(cont_run, qn(W_NS, "continuationSeparator"))

    footnote = ET.SubElement(root, qn(W_NS, "footnote"), {qn(W_NS, "id"): "2"})
    para = ET.SubElement(footnote, qn(W_NS, "p"))
    for child in wrap_with_comment(
        4,
        [
            run_el("Audit reports must be delivered within "),
            deletion_el("5", revision_id=24, author="Vendor Counsel"),
            insertion_el("3", revision_id=25, author="Vendor Counsel"),
            run_el(" business days of written request."),
        ],
    ):
        para.append(child)

    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def add_relationship(rels_root: ET.Element, rel_type: str, target: str) -> None:
    for rel in rels_root.findall(f"{{{PKG_REL_NS}}}Relationship"):
        if rel.get("Type") == rel_type and rel.get("Target") == target:
            return

    ids = []
    for rel in rels_root.findall(f"{{{PKG_REL_NS}}}Relationship"):
        raw_id = rel.get("Id", "")
        if raw_id.startswith("rId"):
            try:
                ids.append(int(raw_id[3:]))
            except ValueError:
                pass

    ET.SubElement(
        rels_root,
        qn(PKG_REL_NS, "Relationship"),
        {
            "Id": f"rId{max(ids, default=0) + 1}",
            "Type": rel_type,
            "Target": target,
        },
    )


def add_content_type(root: ET.Element, part_name: str, content_type: str) -> None:
    for override in root.findall(f"{{{CONTENT_TYPES_NS}}}Override"):
        if override.get("PartName") == part_name:
            return
    ET.SubElement(
        root,
        qn(CONTENT_TYPES_NS, "Override"),
        {
            "PartName": part_name,
            "ContentType": content_type,
        },
    )


def add_track_revisions(settings_root: ET.Element) -> None:
    if settings_root.find(qn(W_NS, "trackRevisions")) is None:
        settings_root.insert(1, ET.Element(qn(W_NS, "trackRevisions")))


def build_base_docx(output_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.header.paragraphs[0].text = "Vendor Services Addendum - Legal Redline"
    section.footer.paragraphs[0].text = "Confidential - North Harbor Procurement Group"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("VENDOR SERVICES ADDENDUM").bold = True

    buyer = doc.add_paragraph()
    buyer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    buyer.add_run("North Harbor Procurement Group").bold = True

    doc.add_paragraph("Date: April 18, 2026")
    doc.add_paragraph("Vendor: Orion Managed Services LLC")
    doc.add_paragraph("Agreement: Infrastructure Monitoring Subscription Renewal")
    doc.add_paragraph("This addendum captures the final legal review decisions before the renewed agreement is circulated for signature.")
    doc.add_paragraph("[[TERM_EXTENSION]]")
    doc.add_paragraph("[[NOTICE_PERIOD]]")
    doc.add_paragraph("[[SECURITY_REVIEW]]")
    doc.add_paragraph("[[USAGE_LOG_RETENTION]]")
    doc.add_paragraph("[[AUDIT_REPORT_WINDOW]]")
    doc.add_paragraph("Commercial terms under review:")

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Annual Fee", "[[PRICING_HOLD]]"),
        ("Licensed Users", "[[LICENSE_COUNT]]"),
        ("Service Credit Cap", "[[SERVICE_CREDIT_CAP]]"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value

    doc.add_paragraph("Please return a fully signed addendum no later than April 25, 2026.")
    doc.add_paragraph("North Harbor Procurement Group Procurement Operations")
    doc.save(output_path)


def inject_redlines(source_docx: Path, output_docx: Path) -> None:
    with ZipFile(source_docx) as source, ZipFile(output_docx, "w", compression=ZIP_DEFLATED) as target:
        for info in source.infolist():
            if info.filename in {
                "customXml/item1.xml",
                "customXml/itemProps1.xml",
                "customXml/_rels/item1.xml.rels",
            }:
                continue
            payload = source.read(info.filename)
            if info.filename == "word/document.xml":
                root = ET.fromstring(payload)
                for para in root.iter(qn(W_NS, "p")):
                    text = "".join(node.text or "" for node in para.iter(qn(W_NS, "t"))).strip()
                    if text == "[[TERM_EXTENSION]]":
                        replace_paragraph_children(
                            para,
                            wrap_with_comment(
                                0,
                                [
                                    run_el("The renewed service term will continue through April 30, "),
                                    deletion_el("2026", revision_id=10, author="Vendor Counsel"),
                                    insertion_el("2027", revision_id=11, author="Vendor Counsel"),
                                    run_el("."),
                                ],
                            ),
                        )
                    elif text == "[[NOTICE_PERIOD]]":
                        replace_paragraph_children(
                            para,
                            wrap_with_comment(
                                1,
                                [
                                    run_el("Either party may elect not to renew by giving "),
                                    deletion_el("30", revision_id=12, author="Legal Ops"),
                                    insertion_el("45", revision_id=13, author="Legal Ops"),
                                    run_el(" days' written notice before the end of the term."),
                                ],
                            ),
                        )
                    elif text == "[[SECURITY_REVIEW]]":
                        replace_paragraph_children(
                            para,
                            wrap_with_comment(
                                2,
                                [
                                    insertion_el(
                                        "A current security review packet must be returned before the renewal takes effect.",
                                        revision_id=14,
                                        author="Security Counsel",
                                    )
                                ],
                            ),
                        )
                    elif text == "[[USAGE_LOG_RETENTION]]":
                        replace_paragraph_children(
                            para,
                            wrap_with_comment(
                                3,
                                [
                                    run_el("System usage logs must be retained for "),
                                    deletion_el("90", revision_id=15, author="Security Counsel"),
                                    insertion_el("180", revision_id=16, author="Security Counsel"),
                                    run_el(" days following each monthly service period."),
                                ],
                            ),
                        )
                    elif text == "[[AUDIT_REPORT_WINDOW]]":
                        replace_paragraph_children(
                            para,
                            [
                                run_el("Security reporting timing is described in the attached audit note"),
                                footnote_reference_run(2),
                                run_el("."),
                            ],
                        )
                    elif text == "[[PRICING_HOLD]]":
                        replace_paragraph_children(
                            para,
                            wrap_with_comment(
                                5,
                                [
                                    run_el("$"),
                                    deletion_el("248,500", revision_id=17, author="Vendor Counsel"),
                                    insertion_el("255,000", revision_id=18, author="Vendor Counsel"),
                                    run_el(" per year"),
                                ],
                            ),
                        )
                    elif text == "[[LICENSE_COUNT]]":
                        replace_paragraph_children(
                            para,
                            wrap_with_comment(
                                6,
                                [
                                    deletion_el("175", revision_id=19, author="Vendor Counsel"),
                                    insertion_el("200", revision_id=20, author="Vendor Counsel"),
                                    run_el(" seats"),
                                ],
                            ),
                        )
                    elif text == "[[SERVICE_CREDIT_CAP]]":
                        replace_paragraph_children(
                            para,
                            wrap_with_comment(
                                7,
                                [
                                    run_el("$"),
                                    deletion_el("12,000", revision_id=21, author="Vendor Counsel"),
                                    insertion_el("18,000", revision_id=22, author="Vendor Counsel"),
                                    run_el(" per year"),
                                ],
                            ),
                        )
                payload = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif info.filename == "word/_rels/document.xml.rels":
                root = ET.fromstring(payload)
                add_relationship(
                    root,
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
                    "comments.xml",
                )
                add_relationship(
                    root,
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
                    "footnotes.xml",
                )
                payload = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif info.filename == "[Content_Types].xml":
                root = ET.fromstring(payload)
                add_content_type(
                    root,
                    "/word/comments.xml",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
                )
                add_content_type(
                    root,
                    "/word/footnotes.xml",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
                )
                payload = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif info.filename == "word/settings.xml":
                root = ET.fromstring(payload)
                add_track_revisions(root)
                payload = ET.tostring(root, encoding="utf-8", xml_declaration=True)

            target.writestr(info, payload)

        target.writestr("word/comments.xml", make_comments_xml())
        target.writestr("word/footnotes.xml", make_footnotes_xml())
        target.writestr("customXml/item1.xml", make_review_manifest_xml())
        target.writestr("customXml/itemProps1.xml", make_review_manifest_props_xml())
        target.writestr("customXml/_rels/item1.xml.rels", make_review_manifest_rels_xml())


def main() -> None:
    script_dir = Path(__file__).resolve().parent
    environment_root = script_dir.parent
    output_dir = environment_root / "data"
    output_dir.mkdir(parents=True, exist_ok=True)

    base_docx = output_dir / "vendor_addendum_base.docx"
    final_docx = output_dir / "vendor_addendum_redline.docx"
    decisions_json = output_dir / "review_decisions.json"

    decisions = {
        "TERM_EXTENSION": "accept",
        "NOTICE_PERIOD": "accept",
        "SECURITY_REVIEW": "accept",
        "USAGE_LOG_RETENTION": "accept",
        "AUDIT_REPORT_WINDOW": "reject",
        "PRICING_HOLD": "reject",
        "LICENSE_COUNT": "accept",
        "SERVICE_CREDIT_CAP": "reject",
    }

    build_base_docx(base_docx)
    inject_redlines(base_docx, final_docx)
    decisions_json.write_text(json.dumps(decisions, indent=2) + "\n", encoding="utf-8")
    base_docx.unlink()


if __name__ == "__main__":
    main()
